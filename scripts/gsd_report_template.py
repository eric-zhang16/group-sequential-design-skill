#!/usr/bin/env python3
"""
GSD Report Template — generates gsd_report.docx from gsd_results.json.
Pure template: ZERO hardcoded design values. Every number comes from the JSON.

Usage:
    python gsd_report_template.py <output_dir>

The script expects:
    <output_dir>/gsd_results.json
    <output_dir>/multiplicity_diagram.png (optional)

It produces:
    <output_dir>/gsd_report.docx
"""

import json
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Parse CLI ────────────────────────────────────────────────────────────────

out_dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
json_path = os.path.join(out_dir, "gsd_results.json")
diagram_path = os.path.join(out_dir, "multiplicity_diagram.png")
report_path = os.path.join(out_dir, "gsd_report.docx")

with open(json_path) as f:
    r = json.load(f)


# ── Helpers ──────────────────────────────────────────────────────────────────

def pct(x, dec=1):
    """Format a proportion (0-1) or percentage (0-100) as 'XX.X%'."""
    v = x * 100 if x <= 1 else x
    return f"{v:.{dec}f}%"

def fmt(x, dec=3):
    return f"{x:.{dec}f}"

def fmt_mo(x):
    return f"{x:.1f}"

def fmt_hr(x):
    return f"{x:.3f}"

def classify_ia_stringency(hr_at_bound, cum_power_at_ia):
    """Classify IA boundary stringency from design values."""
    if hr_at_bound < 0.70 or cum_power_at_ia < 0.50:
        return "stringent"
    elif hr_at_bound > 0.85 or cum_power_at_ia > 0.80:
        return "lenient"
    else:
        return "moderate"

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
    return t

def j(key, default=""):
    """Safe get from JSON with default."""
    keys = key.split(".")
    v = r
    for k in keys:
        if isinstance(v, dict):
            v = v.get(k, default)
        else:
            return default
    return v


# ── Extract common fields ────────────────────────────────────────────────────

disease = j("disease", j("disease_setting", ""))
endpoints_str = j("endpoints", "")
randomization = j("randomization", "1:1")
alpha = j("alpha", 0.025)
total_N = j("total_N", 0)
N_per_arm = j("N_per_arm", total_N // 2)
enroll_dur = j("enroll_duration", j("enrollment_duration", 0))
study_dur = j("study_duration", 0)
power_target = j("power_target_pct", j("power_target", 90))
if power_target <= 1:
    power_target *= 100
n_sim = j("n_sim", 10000)
min_fup = j("min_followup", j("min_fup", 0))
min_gap = j("min_gap", 0)
max_N = j("max_N_constraint", j("n_max_constraint", ""))
feasible_range = j("feasible_range", "")
eff_spending = j("efficacy_spending", j("spending_eff", ""))
fut_spending = j("futility_spending", "")
fut_type = j("futility_type", "")
dropout = j("dropout_annual", 0)

# Hypotheses — detect multi-hypothesis designs
hypotheses = j("hypotheses", None)
is_multi_hyp = hypotheses is not None and isinstance(hypotheses, (list, dict))


# ── Build document ───────────────────────────────────────────────────────────

doc = Document()

# Title
title_p = doc.add_paragraph()
title_run = title_p.add_run(f"Group Sequential Design Report: {disease}")
title_run.bold = True
title_run.font.size = Pt(16)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Section 1: Design Assumptions
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Design Assumptions", 1)

assumptions_rows = [
    ["Disease / Setting", str(disease)],
    ["Endpoints", str(endpoints_str)],
    ["Randomization", str(randomization)],
    ["Total alpha (one-sided)", str(alpha)],
]

# Add control medians and HRs — handle single and multi-endpoint
for key in ["ctrl_median_os", "ctrl_median_pfs", "ctrl_median_dfs"]:
    val = j(key)
    if val:
        label = key.replace("ctrl_median_", "Control median ").upper()
        assumptions_rows.append([label, f"{val} months"])

for key in ["hr_os", "hr_pfs", "hr_dfs"]:
    val = j(key)
    if val:
        label = key.replace("hr_", "Target HR ").upper()
        assumptions_rows.append([label, str(val)])

assumptions_rows.extend([
    ["Total N", str(total_N)],
    ["N per arm", str(N_per_arm)],
    ["Enrollment duration", f"{enroll_dur} months"],
    ["Efficacy spending", str(eff_spending)],
    ["Futility", f"{fut_type} — {fut_spending}" if fut_spending else str(fut_type)],
    ["Annual dropout", pct(dropout) if dropout <= 1 else f"{dropout}%"],
    ["Min follow-up", f"{min_fup} months"],
    ["Min gap between analyses", f"{min_gap} months"],
])

if max_N:
    assumptions_rows.append(["N constraint", f"< {max_N}" if isinstance(max_N, (int, float)) else str(max_N)])
if feasible_range:
    assumptions_rows.append(["Feasible range", str(feasible_range)])

add_table(doc, ["Parameter", "Value"], assumptions_rows)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Section 2: Interim/Final Analysis Plan
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Interim/Final Analysis Plan", 1)

# Try to build the IA/FA plan from JSON
# Look for analysis plan in various JSON structures
analyses = j("analyses", None)
two_ia = j("two_ia", j("design_2ia", None))

if analyses and isinstance(analyses, list):
    plan_headers = ["Analysis", "Timing (mo)", "Events", "IF", "Trigger"]
    plan_rows = []
    for a in analyses:
        plan_rows.append([
            a.get("name", ""),
            fmt_mo(a.get("time", 0)),
            str(a.get("events", "")),
            pct(a.get("if", a.get("info_frac", 0))) if a.get("if", a.get("info_frac")) else "",
            a.get("trigger", a.get("driver", "")),
        ])
    add_table(doc, plan_headers, plan_rows)
else:
    # Reconstruct from individual fields
    ia_time = j("ia_time", j("ia1_time", 0))
    fa_time = j("fa_time", 0)
    events_ia = j("events_ia", j("ia1_events", 0))
    events_fa = j("events_fa", 0)
    os_if_ia = j("os_if_ia", j("os_info_frac", []))

    plan_text = f"The study includes "
    if two_ia and isinstance(two_ia, dict) and two_ia.get("computed", False):
        ia1_t = two_ia.get("ia1_time", 0)
        ia2_t = two_ia.get("ia2_time", 0)
        fa_t = two_ia.get("fa_time", 0)
        plan_text += (f"two interim analyses and a final analysis. "
                      f"IA1 at month {fmt_mo(ia1_t)}, IA2 at month {fmt_mo(ia2_t)}, "
                      f"FA at month {fmt_mo(fa_t)}.")
    elif ia_time and fa_time:
        plan_text += (f"one interim analysis at month {fmt_mo(ia_time)} and "
                      f"a final analysis at month {fmt_mo(fa_time)}.")
    else:
        plan_text += "analyses as specified in the design."

    add_para(doc, plan_text)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Section 3: Multiplicity Strategy
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Multiplicity Strategy", 1)

alpha_split = j("alpha_split", None)
reallocation = j("alpha_reallocation", j("reallocation", ""))

if is_multi_hyp:
    add_para(doc, f"The design tests multiple hypotheses with total alpha = {alpha} (one-sided). "
                  f"Alpha is allocated across hypotheses as specified in the design assumptions.")
    if reallocation:
        add_para(doc, f"Alpha reallocation: {reallocation}")
elif alpha_split:
    add_para(doc, f"Alpha is split between endpoints: {alpha_split}. Total = {alpha} (one-sided).")
    if reallocation:
        add_para(doc, f"Alpha reallocation: {reallocation}")
else:
    add_para(doc, f"Single-hypothesis design with alpha = {alpha} (one-sided).")

# Insert multiplicity diagram if available
if os.path.exists(diagram_path):
    doc.add_picture(diagram_path, width=Inches(5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Section 4: Efficacy and Futility Boundaries
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Efficacy and Futility Boundaries", 1)


def build_boundary_table(doc, label, alpha_val, z_upper, z_lower, hr_upper, hr_lower,
                         p_upper, cum_power, cum_alpha, events, ifs, n_at,
                         analysis_labels=None, has_futility=True):
    """Build a boundary table for one endpoint/hypothesis."""
    k = len(z_upper) if isinstance(z_upper, list) else 1

    # Single-look endpoint
    if k == 1:
        zu = z_upper if not isinstance(z_upper, list) else z_upper[0]
        pu = p_upper if not isinstance(p_upper, list) else p_upper[0]
        hru = hr_upper if not isinstance(hr_upper, list) else hr_upper[0]
        ev = events if not isinstance(events, list) else events[0]

        heading_text = f"{label}, alpha = {alpha_val}"
        add_heading(doc, heading_text, 2)
        add_para(doc, f"Single-look analysis with {ev} events.")

        headers = ["Analysis", "Events", "Z boundary", "p (1-sided)", "HR at boundary"]
        rows = [["Single look", str(ev), fmt(zu, 4), fmt(pu, 4), fmt_hr(hru)]]
        add_table(doc, headers, rows)
        return

    # Multi-look endpoint
    heading_text = f"{label}, alpha = {alpha_val}"
    add_heading(doc, heading_text, 2)

    if has_futility and z_lower:
        headers = ["Analysis", "IF", "N", "Events", "Z eff.", "Z fut.",
                   "p eff. (1-sided)", "HR eff.", "HR fut.",
                   "Cum. power", "Cum. alpha"]
    else:
        headers = ["Analysis", "IF", "N", "Events", "Z boundary",
                   "p (1-sided)", "HR at boundary", "Cum. power", "Cum. alpha"]

    rows = []
    for i in range(k):
        lbl = analysis_labels[i] if analysis_labels and i < len(analysis_labels) else f"Analysis {i+1}"
        ev = events[i] if isinstance(events, list) and i < len(events) else ""
        if_val = pct(ifs[i]) if isinstance(ifs, list) and i < len(ifs) else ""
        n_val = n_at[i] if isinstance(n_at, list) and i < len(n_at) else str(total_N)
        zu = fmt(z_upper[i], 4) if isinstance(z_upper, list) and i < len(z_upper) else ""
        pu = fmt(p_upper[i], 4) if isinstance(p_upper, list) and i < len(p_upper) else ""
        hru = fmt_hr(hr_upper[i]) if isinstance(hr_upper, list) and i < len(hr_upper) else ""
        cp = pct(cum_power[i]) if isinstance(cum_power, list) and i < len(cum_power) else ""
        ca = fmt(cum_alpha[i], 4) if isinstance(cum_alpha, list) and i < len(cum_alpha) else ""

        if has_futility and z_lower:
            zl = fmt(z_lower[i], 4) if isinstance(z_lower, list) and i < len(z_lower) else ""
            hrl = fmt_hr(hr_lower[i]) if isinstance(hr_lower, list) and i < len(hr_lower) else ""
            rows.append([lbl, if_val, str(n_val), str(ev), zu, zl, pu, hru, hrl, cp, ca])
        else:
            rows.append([lbl, if_val, str(n_val), str(ev), zu, pu, hru, cp, ca])

    add_table(doc, headers, rows)


# Detect endpoints and build tables
# Pattern: multi-hypothesis (H1, H2, H3, H4...)
if is_multi_hyp and isinstance(hypotheses, dict):
    for hname, hdata in hypotheses.items():
        build_boundary_table(
            doc, label=f"{hname}: {hdata.get('endpoint', '')} ({hdata.get('population', '')})",
            alpha_val=hdata.get("alpha_initial", hdata.get("alpha", "")),
            z_upper=hdata.get("z_upper", []),
            z_lower=hdata.get("z_lower", []),
            hr_upper=hdata.get("hr_upper", hdata.get("hr_eff", [])),
            hr_lower=hdata.get("hr_lower", hdata.get("hr_fut", [])),
            p_upper=hdata.get("p_upper", hdata.get("p_eff", [])),
            cum_power=hdata.get("cum_power", []),
            cum_alpha=hdata.get("cum_alpha", []),
            events=hdata.get("events", []),
            ifs=hdata.get("info_frac", hdata.get("if", [])),
            n_at=hdata.get("n_at", []),
            has_futility=hdata.get("has_futility", False),
        )
else:
    # Single or co-primary — look for os_*, pfs_* prefixed fields
    for prefix, ep_label in [("pfs", "PFS"), ("os", "OS"), ("dfs", "DFS")]:
        z_up = j(f"{prefix}_z_upper")
        if not z_up:
            continue
        build_boundary_table(
            doc, label=ep_label,
            alpha_val=j(f"alpha_{prefix}", j(f"{prefix}_alpha", alpha)),
            z_upper=z_up,
            z_lower=j(f"{prefix}_z_lower"),
            hr_upper=j(f"{prefix}_hr_upper", j(f"{prefix}_hr_eff")),
            hr_lower=j(f"{prefix}_hr_lower", j(f"{prefix}_hr_fut")),
            p_upper=j(f"{prefix}_p_upper", j(f"{prefix}_p_eff")),
            cum_power=j(f"{prefix}_cum_cross_h1", j(f"{prefix}_cum_power")),
            cum_alpha=j(f"{prefix}_cum_alpha"),
            events=[j("events_ia", j(f"{prefix}_events_ia")), j("events_fa", j(f"{prefix}_events_fa"))],
            ifs=j(f"{prefix}_info_frac", j(f"{prefix}_if")),
            n_at=j(f"{prefix}_N_at_analysis", [total_N]),
            has_futility=(prefix == "os" and fut_type),
        )

    # If no prefix-based fields found, try generic single-endpoint
    if not any(j(f"{p}_z_upper") for p in ["pfs", "os", "dfs"]):
        z_up = j("os_z_upper", j("z_upper"))
        if z_up:
            build_boundary_table(
                doc, label=endpoints_str or "OS",
                alpha_val=alpha,
                z_upper=z_up,
                z_lower=j("os_z_lower", j("z_lower")),
                hr_upper=j("os_hr_upper", j("hr_upper", j("os_hr_eff"))),
                hr_lower=j("os_hr_lower", j("hr_lower", j("os_hr_fut"))),
                p_upper=j("os_p_upper", j("p_upper", j("os_p_eff"))),
                cum_power=j("os_cum_cross_h1", j("cum_power")),
                cum_alpha=j("os_cum_alpha", j("cum_alpha")),
                events=[j("events_ia"), j("events_fa")],
                ifs=j("os_info_frac", j("info_frac")),
                n_at=j("os_N_at_analysis", [total_N]),
                has_futility=bool(fut_type),
            )

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Section 5: Sample Size and Power Summary
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Sample Size and Power Summary", 1)

add_para(doc, f"The design enrolls {total_N} patients ({N_per_arm} per arm) "
              f"over {enroll_dur} months with a total study duration of "
              f"{fmt_mo(study_dur)} months.")

# Power summary
os_power = j("os_power_pct", j("os_power", j("power_ph", 0)))
if os_power and os_power <= 1:
    os_power *= 100
pfs_power = j("pfs_power_pct", j("pfs_power", 0))
if pfs_power and pfs_power <= 1:
    pfs_power *= 100

power_rows = []
if os_power:
    power_rows.append(["OS", f"{os_power:.1f}%"])
if pfs_power:
    power_rows.append(["PFS", f"{pfs_power:.1f}%"])

if power_rows:
    add_table(doc, ["Endpoint", "Power"], power_rows)

# NPH power if present
nph_power = j("nph_power_fa", j("nph_power_total", 0))
if nph_power:
    if nph_power <= 1:
        nph_power *= 100
    doc.add_paragraph()
    add_para(doc, f"Under non-proportional hazards assumptions, "
                  f"power at the final analysis is {nph_power:.1f}%.")

    # AHR table
    ahr_data = []
    for suffix in ["ia1", "ia2", "fa"]:
        ahr = j(f"nph_ahr_{suffix}")
        if ahr:
            label = {"ia1": "IA1", "ia2": "IA2", "fa": "FA"}[suffix]
            ahr_data.append([label, fmt(ahr, 4)])
    if ahr_data:
        add_para(doc, "Average Hazard Ratio (AHR) under NPH:", bold=True)
        add_table(doc, ["Analysis", "AHR"], ahr_data)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Section 6: Design Assessment
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Design Assessment", 1)

# --- Strengths ---
add_heading(doc, "Strengths", 2)
strengths = []

if os_power and os_power >= power_target:
    strengths.append(f"Achieved {os_power:.1f}% power with {total_N} patients, "
                     f"meeting the {power_target:.0f}% target.")

if "OBF" in str(eff_spending) or "sfLDOF" in str(eff_spending):
    cum_alpha_ia = j("os_cum_alpha", j("cum_alpha", []))
    if isinstance(cum_alpha_ia, list) and len(cum_alpha_ia) > 0:
        strengths.append(f"Lan-DeMets O'Brien-Fleming spending preserves alpha for the final analysis "
                         f"(cumulative alpha at first IA = {cum_alpha_ia[0]:.4f}).")

if "non-binding" in str(fut_type).lower() or "non_binding" in str(fut_type).lower():
    strengths.append("Non-binding futility provides advisory guidance without mandating early "
                     "termination and does not inflate the type I error.")

for s in strengths:
    add_para(doc, f"  \u2022  {s}")

# --- Limitations ---
add_heading(doc, "Limitations and Considerations", 2)
limitations = []

if study_dur and study_dur > 48:
    limitations.append(f"Study duration of {fmt_mo(study_dur)} months ({study_dur/12:.1f} years) "
                       f"is lengthy, driven by the need to accumulate sufficient events in {total_N} patients.")

# Gap check
gap = j("gap_months", j("gap_mo", 0))
gap_threshold = j("gap_warn_threshold", 18)
if gap and gap > gap_threshold:
    limitations.append(f"The IA-to-FA gap of {fmt_mo(gap)} months exceeds the "
                       f"{gap_threshold}-month threshold.")

# Futility conservatism
hr_fut_ia1 = None
os_hr_lower = j("os_hr_lower", j("os_hr_fut", []))
if isinstance(os_hr_lower, list) and len(os_hr_lower) > 0:
    hr_fut_ia1 = os_hr_lower[0]
elif isinstance(os_hr_lower, (int, float)):
    hr_fut_ia1 = os_hr_lower

if hr_fut_ia1 and hr_fut_ia1 > 0.90:
    hr_design = j("hr_os", j("hr", 0))
    limitations.append(f"Futility boundary at first IA (HR threshold = {fmt_hr(hr_fut_ia1)}) "
                       f"is very conservative — nearly at the null. Under the design alternative "
                       f"(HR = {hr_design}), this boundary is rarely crossed.")

# IA stringency
os_hr_upper = j("os_hr_upper", j("os_hr_eff", []))
os_cum_power = j("os_cum_cross_h1", j("os_cum_power", []))
if isinstance(os_hr_upper, list) and isinstance(os_cum_power, list):
    for i in range(len(os_hr_upper) - 1):  # Skip FA
        sev = classify_ia_stringency(os_hr_upper[i], os_cum_power[i])
        if sev == "stringent":
            limitations.append(f"IA{i+1} efficacy boundary is stringent "
                               f"(HR = {fmt_hr(os_hr_upper[i])}, cumulative power = {pct(os_cum_power[i])}). "
                               f"A strong early signal is needed to stop.")

for lim in limitations:
    add_para(doc, f"  \u2022  {lim}")

# --- Potential Improvements ---
add_heading(doc, "Potential Improvements", 2)
improvements = []

# Alpha reallocation if overpowered
if pfs_power and pfs_power > power_target + 3:
    pfs_alpha = j("alpha_pfs", j("pfs_alpha", 0))
    os_alpha = j("alpha_os", j("os_alpha", 0))
    if pfs_alpha and os_alpha:
        improvements.append(f"PFS is overpowered ({pfs_power:.1f}% vs {power_target:.0f}% target). "
                            f"Consider reducing PFS alpha from {pfs_alpha} and increasing OS alpha from {os_alpha} "
                            f"to improve OS power.")

# N increase for shorter study
if study_dur and study_dur > 48:
    improvements.append(f"Increasing enrollment (extending steady-state period) would accelerate "
                        f"event accrual and shorten the study duration, at the cost of a larger trial.")

# Less conservative futility
if hr_fut_ia1 and hr_fut_ia1 > 0.90:
    improvements.append("A less conservative futility spending function (e.g., HSD gamma=-2) "
                        "would produce actionable futility boundaries closer to the design HR.")

# 2-IA suggestion if gap flagged
if two_ia and isinstance(two_ia, dict) and two_ia.get("computed", False):
    ia2_fa_gap = two_ia.get("gap_ia2_fa", two_ia.get("gap_2f", 0))
    if ia2_fa_gap and ia2_fa_gap > gap_threshold:
        improvements.append(f"Even with the added interim, the last gap ({fmt_mo(ia2_fa_gap)} months) "
                            f"exceeds {gap_threshold} months. Consider a third interim at ~"
                            f"{int((two_ia.get('ia2_if', 0.85) * 100 + 100) / 2)}% of events.")

for imp in improvements:
    add_para(doc, f"  \u2022  {imp}")

if not improvements:
    add_para(doc, "No immediate improvements identified. The design meets all targets.")

doc.add_paragraph()


# ── Verification summary (if present) ────────────────────────────────────────

verification = j("verification", None)
if verification and isinstance(verification, dict):
    add_heading(doc, "Appendix: Verification Summary", 1)

    for sim_key in ["sim_1ia", "sim_2ia", "sim_ph", "sim_h1"]:
        sim = verification.get(sim_key)
        if not sim:
            continue
        label = sim_key.replace("sim_", "").upper()
        add_para(doc, f"Verification ({label}):", bold=True)
        add_para(doc, f"  Simulated power: {sim.get('power_sim', 'N/A')}%")
        add_para(doc, f"  Simulated type I error: {sim.get('t1err_sim', 'N/A')}%")
        add_para(doc, f"  Simulations: {n_sim}")
        overall = "PASS" if sim.get("pass", False) else "FAIL"
        add_para(doc, f"  Overall: {overall}", bold=True)


# ── Save ─────────────────────────────────────────────────────────────────────

doc.save(report_path)
print(f"Saved: {report_path}")
